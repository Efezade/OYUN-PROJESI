# -*- coding: utf-8 -*-
"""Karo_Tasarim_Klavuzu.md -> .docx (python-docx). Basit ama saglam MD donusturucu."""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\3D OYUN\OYUN\Docs\Karo_Tasarim_Klavuzu.md"
import sys
OUT = sys.argv[1] if len(sys.argv) > 1 else r"C:\3D OYUN\OYUN\Docs\Karo_Tasarim_Klavuzu.docx"

doc = Document()
# Varsayilan font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def add_inline(paragraph, text):
    """**bold** ve `code` destekli inline parser."""
    # Token: bold veya code veya normal
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = paragraph.add_run(p[2:-2])
            r.bold = True
        elif p.startswith("`") and p.endswith("`"):
            r = paragraph.add_run(p[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(p)

with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    stripped = line.strip()

    # Kod blogu
    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < n and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        i += 1  # kapanis ```
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Pt(6)
        pf.space_before = Pt(4)
        pf.space_after = Pt(8)
        r = p.add_run("\n".join(code_lines))
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        # gri arka plan (shading)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        continue

    # Tablo
    if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|?\s*$", lines[i+1].strip()):
        # tablo satirlarini topla
        tbl_rows = []
        while i < n and lines[i].strip().startswith("|"):
            tbl_rows.append(lines[i].strip())
            i += 1
        # ayrac satirini (---) at
        header = [c.strip() for c in tbl_rows[0].strip("|").split("|")]
        body = []
        for row in tbl_rows[2:]:
            cells = [c.strip() for c in row.strip("|").split("|")]
            body.append(cells)
        ncol = len(header)
        table = doc.add_table(rows=1, cols=ncol)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for j, h in enumerate(header):
            hdr[j].paragraphs[0].text = ""
            add_inline(hdr[j].paragraphs[0], h)
            for run in hdr[j].paragraphs[0].runs:
                run.bold = True
            shade_cell(hdr[j], "D9E2F3")
        for cells in body:
            row_cells = table.add_row().cells
            for j in range(ncol):
                txt = cells[j] if j < len(cells) else ""
                row_cells[j].paragraphs[0].text = ""
                add_inline(row_cells[j].paragraphs[0], txt)
        doc.add_paragraph()
        continue

    # Basliklar
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        i += 1; continue
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        i += 1; continue
    if stripped.startswith("# "):
        h = doc.add_heading(stripped[2:], level=0)
        i += 1; continue

    # Yatay cizgi
    if stripped == "---":
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pbdr.append(bottom)
        pPr.append(pbdr)
        i += 1; continue

    # Blockquote
    if stripped.startswith(">"):
        quote_text = []
        while i < n and lines[i].strip().startswith(">"):
            quote_text.append(lines[i].strip().lstrip(">").strip())
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        add_inline(p, " ".join(quote_text))
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        continue

    # Liste / checkbox
    m = re.match(r"^(\s*)-\s+(.*)$", line)
    if m:
        content = m.group(2)
        content = re.sub(r"^\[\s\]\s*", "☐ ", content)  # bos checkbox
        content = re.sub(r"^\[x\]\s*", "☑ ", content, flags=re.I)
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, content)
        i += 1; continue

    # Bos satir
    if stripped == "":
        i += 1; continue

    # Normal paragraf
    p = doc.add_paragraph()
    add_inline(p, stripped)
    i += 1

doc.save(OUT)
print("OK ->", OUT)
